import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def parse_markdown(filepath):
    with open(filepath, 'r', encoding='utf-8') as f:
        content = f.read()
    
    elements = []
    lines = content.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]
        
        # 标题
        if line.startswith('# '):
            elements.append(('h1', line[2:].strip()))
        elif line.startswith('## '):
            elements.append(('h2', line[3:].strip()))
        elif line.startswith('### '):
            elements.append(('h3', line[4:].strip()))
        elif line.startswith('#### '):
            elements.append(('h4', line[5:].strip()))
        # 表格
        elif '|' in line and i + 1 < len(lines) and '|---' in lines[i + 1]:
            # 收集表格所有行
            table_lines = [line]
            i += 1
            while i < len(lines) and '|' in lines[i]:
                table_lines.append(lines[i])
                i += 1
            elements.append(('table', table_lines))
            continue
        # 分隔线
        elif line.strip() == '---':
            elements.append(('hr', ''))
        # 空行
        elif line.strip() == '':
            elements.append(('empty', ''))
        # 普通段落
        else:
            # 收集连续的非空行作为同一段落
            para_lines = [line]
            i += 1
            while i < len(lines) and lines[i].strip() != '' and not lines[i].startswith('# ') and not lines[i].startswith('## ') and not lines[i].startswith('### ') and not lines[i].startswith('#### ') and not (lines[i].strip() == '---') and not ('|' in lines[i] and i + 1 < len(lines) and '|---' in lines[i + 1]):
                para_lines.append(lines[i])
                i += 1
            elements.append(('p', '\n'.join(para_lines)))
            continue
        
        i += 1
    
    return elements

def format_text(run, text):
    """处理文本中的加粗和斜体"""
    # 处理加粗 **text**
    bold_pattern = r'\*\*(.+?)\*\*'
    parts = re.split(bold_pattern, text)
    
    first_run = True
    for part in parts:
        if part.startswith('**') or part.endswith('**'):
            continue
        if re.match(bold_pattern, '**' + part + '**'):
            r = run if first_run else run._element.makeelement(run.tag, run.nsmap)
            r.text = part
            r.bold = True
            first_run = False
        else:
            if first_run:
                run.text = part
                first_run = False
            else:
                r = run._element.makeelement(run.tag, run.nsmap)
                r.text = part
    
    if first_run:
        run.text = text

def create_docx(elements, output_path):
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    # 设置段落间距
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.line_spacing = 1.5
    
    for elem_type, elem_content in elements:
        if elem_type == 'h1':
            p = doc.add_heading(elem_content, level=1)
            for run in p.runs:
                run.font.name = '黑体'
                run.font.color.rgb = RGBColor(0, 0, 0)
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(24)
            p.paragraph_format.space_after = Pt(12)
        
        elif elem_type == 'h2':
            p = doc.add_heading(elem_content, level=2)
            for run in p.runs:
                run.font.name = '黑体'
                run.font.color.rgb = RGBColor(0, 0, 0)
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(10)
        
        elif elem_type == 'h3':
            p = doc.add_heading(elem_content, level=3)
            for run in p.runs:
                run.font.name = '黑体'
                run.font.color.rgb = RGBColor(0, 0, 0)
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
        
        elif elem_type == 'h4':
            p = doc.add_heading(elem_content, level=4)
            for run in p.runs:
                run.font.name = '黑体'
                run.font.color.rgb = RGBColor(0, 0, 0)
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
        
        elif elem_type == 'p':
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Inches(0.3)
            
            text = elem_content
            # 处理加粗
            bold_pattern = r'\*\*(.+?)\*\*'
            parts = re.split(bold_pattern, text)
            
            for part in parts:
                if part:
                    is_bold = bool(re.match(bold_pattern, '**' + part + '**'))
                    run = p.add_run(part)
                    run.font.size = Pt(12)
                    run.bold = is_bold
                    if is_bold:
                        run.font.name = '黑体'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                    else:
                        run.font.name = '宋体'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        elif elem_type == 'table':
            table_lines = elem_content
            if len(table_lines) >= 2:
                # 解析表格
                headers = [cell.strip() for cell in table_lines[0].split('|')[1:-1]]
                data_rows = []
                for line in table_lines[2:]:  # 跳过分隔行
                    if '|' in line:
                        row = [cell.strip() for cell in line.split('|')[1:-1]]
                        data_rows.append(row)
                
                # 创建表格
                table = doc.add_table(rows=len(data_rows) + 1, cols=len(headers))
                table.style = 'Table Grid'
                
                # 表头
                for j, header in enumerate(headers):
                    cell = table.cell(0, j)
                    cell.text = header
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in paragraph.runs:
                            run.bold = True
                            run.font.name = '黑体'
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                
                # 数据行
                for i, row in enumerate(data_rows):
                    for j, cell_text in enumerate(row):
                        if i < len(data_rows) and j < len(headers):
                            cell = table.cell(i + 1, j)
                            cell.text = cell_text
                            for paragraph in cell.paragraphs:
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                for run in paragraph.runs:
                                    run.font.name = '宋体'
                                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            
            doc.add_paragraph()  # 表格后空一行
        
        elif elem_type == 'hr':
            doc.add_paragraph()  # 空行代替分隔线
    
    doc.save(output_path)
    print(f"成功生成Word文档: {output_path}")

if __name__ == '__main__':
    md_path = r'k:\桌面\毕设\PoemKBQA-master\PoemWebApp\thesis.md'
    output_path = r'k:\桌面\毕设\PoemKBQA-master\PoemWebApp\thesis.docx'
    
    elements = parse_markdown(md_path)
    create_docx(elements, output_path)
